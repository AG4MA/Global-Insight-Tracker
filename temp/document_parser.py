# -*- coding: utf-8 -*-
"""
Document Parser - Estrae testo e metadata da PDF e DOCX
Prepara documenti per analisi AI

Autore: Senior Python Developer
Data: 22 Dicembre 2025
"""

import os
from typing import Dict, Optional, List
from pathlib import Path
import re

# PDF parsing
import PyPDF2
import pdfplumber

# DOCX parsing
from docx import Document as DocxDocument

import utils


class DocumentParser:
    """Parser per estrarre testo da documenti"""
    
    def __init__(self):
        self.logger = utils.logger
    
    def parse_document(self, filepath: str) -> Dict:
        """
        Parse documento e estrai contenuto
        
        Args:
            filepath: Path al documento
        
        Returns:
            Dict con:
            - text: Testo estratto
            - pages: Numero pagine
            - metadata: Metadata estratti
            - tables: Tabelle estratte (se presenti)
        """
        filepath = Path(filepath)
        
        if not filepath.exists():
            self.logger.error(f"❌ File non trovato: {filepath}")
            return None
        
        extension = filepath.suffix.lower()
        
        if extension == '.pdf':
            return self._parse_pdf(filepath)
        elif extension in ['.docx', '.doc']:
            return self._parse_docx(filepath)
        else:
            self.logger.warning(f"⚠️  Formato non supportato: {extension}")
            return None
    
    def _parse_pdf(self, filepath: Path) -> Dict:
        """Parse PDF usando pdfplumber e PyPDF2"""
        
        self.logger.info(f"📄 Parsing PDF: {filepath.name}")
        
        result = {
            'filepath': str(filepath),
            'filename': filepath.name,
            'text': '',
            'pages': 0,
            'metadata': {},
            'tables': [],
            'parse_method': 'pdf'
        }
        
        try:
            # Prima prova con pdfplumber (migliore per testo e tabelle)
            with pdfplumber.open(filepath) as pdf:
                result['pages'] = len(pdf.pages)
                
                # Estrai metadata
                if pdf.metadata:
                    result['metadata'] = {
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'keywords': pdf.metadata.get('Keywords', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                        'producer': pdf.metadata.get('Producer', ''),
                        'creation_date': pdf.metadata.get('CreationDate', '')
                    }
                
                # Estrai testo da ogni pagina
                text_parts = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Estrai tabelle
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            result['tables'].append({
                                'page': i + 1,
                                'data': table
                            })
                
                result['text'] = '\n\n'.join(text_parts)
            
            # Se pdfplumber fallisce, fallback a PyPDF2
            if not result['text']:
                self.logger.info("  Tentativo con PyPDF2...")
                with open(filepath, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    result['pages'] = len(pdf_reader.pages)
                    
                    text_parts = []
                    for page in pdf_reader.pages:
                        text_parts.append(page.extract_text())
                    
                    result['text'] = '\n\n'.join(text_parts)
            
            # Pulisci testo
            result['text'] = self._clean_text(result['text'])
            
            # Statistiche
            word_count = len(result['text'].split())
            char_count = len(result['text'])
            
            self.logger.info(f"  ✅ Estratto: {result['pages']} pagine, "
                           f"{word_count} parole, {char_count} caratteri")
            
            return result
        
        except Exception as e:
            self.logger.error(f"❌ Errore parsing PDF {filepath.name}: {e}")
            return None
    
    def _parse_docx(self, filepath: Path) -> Dict:
        """Parse DOCX"""
        
        self.logger.info(f"📄 Parsing DOCX: {filepath.name}")
        
        result = {
            'filepath': str(filepath),
            'filename': filepath.name,
            'text': '',
            'pages': 0,  # DOCX non ha concetto di pagine
            'metadata': {},
            'tables': [],
            'parse_method': 'docx'
        }
        
        try:
            doc = DocxDocument(filepath)
            
            # Estrai metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }
            
            # Estrai testo dai paragrafi
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            result['text'] = '\n\n'.join(text_parts)
            
            # Estrai tabelle
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                result['tables'].append({
                    'table_number': i + 1,
                    'data': table_data
                })
            
            # Pulisci testo
            result['text'] = self._clean_text(result['text'])
            
            word_count = len(result['text'].split())
            char_count = len(result['text'])
            
            self.logger.info(f"  ✅ Estratto: {word_count} parole, "
                           f"{len(doc.paragraphs)} paragrafi, "
                           f"{len(doc.tables)} tabelle")
            
            return result
        
        except Exception as e:
            self.logger.error(f"❌ Errore parsing DOCX {filepath.name}: {e}")
            return None
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Pulisce testo estratto"""
        
        # Rimuovi caratteri di controllo
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalizza spazi bianchi
        text = re.sub(r'\s+', ' ', text)
        
        # Normalizza newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()
    
    def extract_key_sections(self, parsed_doc: Dict) -> Dict[str, str]:
        """
        Estrae sezioni chiave da un documento (Executive Summary, Conclusions, etc.)
        
        Args:
            parsed_doc: Documento già parsato
        
        Returns:
            Dict con sezioni estratte
        """
        text = parsed_doc.get('text', '')
        
        sections = {}
        
        # Pattern per trovare sezioni comuni
        section_patterns = {
            'executive_summary': [
                r'executive summary',
                r'summary',
                r'overview'
            ],
            'introduction': [
                r'introduction',
                r'background'
            ],
            'key_findings': [
                r'key findings',
                r'main findings',
                r'key insights',
                r'highlights'
            ],
            'conclusions': [
                r'conclusions?',
                r'final thoughts',
                r'closing remarks'
            ],
            'recommendations': [
                r'recommendations?',
                r'next steps',
                r'action items'
            ]
        }
        
        text_lower = text.lower()
        
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                # Cerca il pattern
                matches = list(re.finditer(rf'\b{pattern}\b', text_lower))
                
                if matches:
                    # Prendi la prima occorrenza
                    start = matches[0].start()
                    
                    # Trova la fine (prossima sezione o dopo N caratteri)
                    end = start + 2000  # Max 2000 caratteri per sezione
                    
                    # Cerca se c'è un'altra sezione dopo
                    next_section = re.search(r'\n\s*[A-Z][A-Za-z\s]+\n', 
                                            text[start+50:start+2000])
                    if next_section:
                        end = start + 50 + next_section.start()
                    
                    section_text = text[start:end].strip()
                    sections[section_name] = section_text
                    break
        
        return sections
